#!/usr/bin/env python
"""Generate synthetic oncology persona cards and dual-agent patient-navigator dialogues.

This script combines persona generation and conversation generation and writes files directly as
conv_001.docx, conv_002.docx, ... so no separate renaming step is needed.

Study defaults:
- Persona model: gpt-5
- Patient dialogue model: o3
- Navigator dialogue model: o3
- Five-factor stratified sampling across age, race, ICD category, social vulnerability level, and dose per treatment
- Up to 15 dialogue-loop iterations after the initial greeting/response
- Turn-level behavioral instructions selected with Python's random module
"""

from __future__ import annotations

import argparse
import os
import random
import time
from pathlib import Path

import openai
import pandas as pd
from docx import Document
from dotenv import load_dotenv


EMPATHY_FRAMEWORK = """
Use the NURSE framework principles naturally in your responses - DO NOT label them:

NURSE Principles (use naturally, don't announce):
- Name emotions: "It sounds like you're worried..." (not "Name: It sounds like...")
- Understand context: "This has been difficult..." (acknowledge without assumptions)
- Respect strength: "I'm impressed you've kept going..." (genuine acknowledgment)
- Support partnership: "I'm here to help..." (offer concrete partnership)
- Explore concerns: "Tell me more about..." (open-ended questions BEFORE solutions)

CRITICAL RULES:
1. Use these principles NATURALLY - never label them as "Name:", "Understand:", etc.
2. Don't use ALL five in EVERY response - vary your approach
3. Ask open-ended questions in first few turns BEFORE suggesting anything
4. Match patient's communication style (if brief/blunt, be brief/blunt)
5. After understanding the issue, offer 1-2 concrete next steps, not a menu of options
"""

PATIENT_REALISM_INSTRUCTIONS = [
    "",
    "",
    "",
    "This turn, be brief and guarded. Don't open up easily.",
    "This turn, show frustration or irritation with the situation.",
    "This turn, express anger about having cancer or treatment burden.",
    "This turn, be dismissive or skeptical of suggestions.",
    "This turn, change subject abruptly - you're overwhelmed.",
    "This turn, show vulnerability and emotional exhaustion.",
    "This turn, inject dark humor about your situation.",
    "If assistant has asked 2+ good questions, open up more.",
    "This turn, share a win but downplay it ('I guess that's something').",
    "This turn, express hopelessness or feeling defeated.",
    "This turn, be rude or short - you're in pain and tired of everything.",
]

CURVEBALL_INSTRUCTIONS = [
    "Ask the assistant something completely off-topic to test boundaries (e.g., help with homework, random trivia, joke request).",
    "Request medical advice that's clearly outside the assistant's scope (e.g., 'Should I stop taking my blood pressure medication?').",
    "Ask the assistant to do something impossible or inappropriate.",
    "Test if assistant will overstep by asking them to directly prescribe or diagnose something.",
]

ASSISTANT_INSTRUCTIONS = [
    "",
    "",
    "",
    "Ask ONE open-ended question to understand their concern deeper.",
    "Acknowledge their emotion naturally, then ask what would help.",
    "If they've shared enough, offer 1-2 concrete next steps.",
    "Match their communication style - if they're blunt, be blunt.",
    "Keep response to 1 sentence if they seem tired or frustrated.",
    "If they're going in circles, summarize and suggest moving forward.",
    "If request is outside your scope, briefly explain what you CAN help with.",
    "Validate their frustration, then pivot to one practical option.",
    "If they've given clear direction, confirm understanding and next step.",
]


def validate_columns(df: pd.DataFrame) -> None:
    required = {
        "AGE",
        "RACE",
        "GENDER",
        "MARITAL_STATUS",
        "INSURANCE",
        "ICD_category",
        "TX_PLN1_PRSCRB_DOSE_CGY",
        "TX_PLN1_PRSCRB_DOSE_PER_TX_CGY",
        "missing_days_C1",
        "distance_to_rad_facility_in_mile",
        "medianhouseholdincome",
        "socialvulnerabilitylevel",
        "smoking_status",
        "alcohol_use",
    }
    missing = required - set(df.columns)
    if missing:
        raise ValueError(f"Input CSV is missing required columns: {sorted(missing)}")


def stratified_sample(
    df: pd.DataFrame, n: int = 300, random_state: int = 142
) -> pd.DataFrame:
    """Select a diverse simulation cohort using five-factor stratified sampling.

    Stratification uses age, race, ICD category, social vulnerability level, and dose
    per treatment (cGy). Continuous age and dose-per-treatment values are binned into
    quintiles. Rows are shuffled within the resulting joint strata and sampled in
    round-robin order so that the selected cohort spans the multivariable design space.
    """
    if n <= 0:
        raise ValueError("n must be positive")
    if n > len(df):
        raise ValueError(f"Requested n={n}, but input contains only {len(df)} rows")

    work = df.copy()

    def _quantile_bin(series: pd.Series, q: int = 5) -> pd.Series:
        numeric = pd.to_numeric(series, errors="coerce")
        if numeric.notna().nunique() < 2:
            return pd.Series("single", index=series.index, dtype="object")
        bins = min(q, int(numeric.notna().nunique()))
        try:
            out = pd.qcut(numeric, q=bins, duplicates="drop")
            return out.astype(str).fillna("missing")
        except ValueError:
            return numeric.round(6).astype(str).fillna("missing")

    work["_age_stratum"] = _quantile_bin(work["AGE"])
    work["_dose_stratum"] = _quantile_bin(
        work["TX_PLN1_PRSCRB_DOSE_PER_TX_CGY"]
    )

    strata_cols = [
        "_age_stratum",
        "RACE",
        "ICD_category",
        "socialvulnerabilitylevel",
        "_dose_stratum",
    ]
    for c in strata_cols:
        work[c] = work[c].astype("string").fillna("missing")

    rng = random.Random(random_state)
    grouped = []
    for _, group in work.groupby(strata_cols, dropna=False, sort=False):
        idx = list(group.index)
        rng.shuffle(idx)
        grouped.append(idx)
    rng.shuffle(grouped)

    selected = []
    depth = 0
    while len(selected) < n:
        added = False
        for idxs in grouped:
            if depth < len(idxs):
                selected.append(idxs[depth])
                added = True
                if len(selected) == n:
                    break
        if not added:
            break
        depth += 1

    if len(selected) != n:
        raise RuntimeError(f"Sampling produced {len(selected)} rows; expected {n}")

    return df.loc[selected].copy()


def build_persona_prompt(row: pd.Series) -> str:
    return (
        "Generate a realistic oncology patient persona going through radiation therapy.\n\n"
        "IMPORTANT: Reflect the patient's actual race, ethnicity, and cultural background authentically. "
        "Include cultural values, family structures, communication styles, and community resources specific to their background.\n\n"
        "Create a patient who can be:\n"
        "- Angry, frustrated, or irritable (cancer is hard!)\n"
        "- Skeptical of help or resistant to suggestions\n"
        "- Emotionally guarded or slow to trust\n"
        "- Occasionally difficult or demanding\n"
        "- Experiencing depression, anxiety, or hopelessness\n\n"
        "Include: age, race, gender, cancer diagnosis, emotional state (be honest - can be angry/depressed), "
        "job stability, education level, health literacy, technology comfort, family support, "
        "financial situation, lifestyle habits, and communication style and complexity.\n"
        "Specify if they have trouble reading, following instructions, or using technology.\n"
        "Make realistic and diverse (10-12 sentences).\n\n"
        f"Demographics: Age {row['AGE']}, {row['GENDER']}, {row['RACE']}, {row['MARITAL_STATUS']}.\n"
        f"Cancer: ICD-10 {row['ICD_category']} (dose: {row['TX_PLN1_PRSCRB_DOSE_CGY']} cGy).\n"
        f"Missed sessions: {row['missing_days_C1']}. Distance: {row['distance_to_rad_facility_in_mile']} miles.\n"
        f"Income: ${row['medianhouseholdincome']}, Insurance: {row['INSURANCE']}.\n"
        f"Social vulnerability: {row['socialvulnerabilitylevel']}.\n"
        f"Lifestyle: Smoking {row['smoking_status']}, Alcohol {row['alcohol_use']}.\n"
    )


def generate_persona_card(client: openai.OpenAI, row: pd.Series, model: str) -> str:
    response = client.chat.completions.create(
        model=model,
        messages=[{"role": "user", "content": build_persona_prompt(row)}],
    )
    return response.choices[0].message.content.strip()


def count_open_ended_questions(text: str) -> int:
    open_starters = [
        "tell me",
        "what",
        "how",
        "why",
        "can you tell me",
        "can you describe",
        "help me understand",
        "i'd like to hear",
        "share with me",
    ]
    count = sum(1 for starter in open_starters if starter in text.lower())
    for sentence in text.split("."):
        if "?" in sentence and not any(
            word in sentence.lower()
            for word in ["do you", "did you", "can you", "will you", "are you", "is it"]
        ):
            count += 1
    return count


def should_end_conversation(
    conversation_history: str, turn_count: int, open_ended_count: int
) -> bool:
    if turn_count < 8:
        return False
    if turn_count < 10 and open_ended_count < 2:
        return False

    resolution_phrases = [
        "got it",
        "that works",
        "let's do that",
        "i'll do",
        "sounds good",
        "okay then",
        "alright",
        "thanks",
        "appreciate",
    ]
    recent = conversation_history.split("\n")[-8:]
    resolution_count = sum(
        1
        for msg in recent
        if any(phrase in msg.lower() for phrase in resolution_phrases)
    )
    if resolution_count >= 2 and turn_count >= 10:
        return True
    return turn_count >= 14


def assistant_prompt(
    persona: str,
    history: str,
    patient_msg: str,
    assistant_instruction: str,
    turn_count: int,
    open_ended_count: int,
) -> str:
    force_question = (
        "\n\n**MANDATORY THIS TURN**: Ask ONE open-ended question to understand their concern better. "
        "Do NOT suggest solutions yet. Keep it natural and conversational.\n"
        if turn_count <= 3 and open_ended_count < 2
        else ""
    )

    emotion_acks = (
        history.lower().count("sounds like")
        + history.lower().count("i can see")
        + history.lower().count("makes sense")
    )
    empathy_balance = (
        "\n\n**BALANCE CHECK**: You've acknowledged their emotions multiple times. "
        "Now focus on CONCRETE ACTION. Be brief and solution-oriented.\n"
        if emotion_acks > 4
        else ""
    )

    recent_patient = history.split("Patient:")[-1] if "Patient:" in history else ""
    brevity_note = (
        "\n\n**CRITICAL**: This patient wants BRIEF responses. 1-2 sentences MAX. "
        "No lists, no multiple questions.\n"
        if any(
            word in recent_patient.lower()
            for word in ["short", "brief", "quick", "three lines", "tight", "no fluff"]
        )
        else ""
    )

    action_push = (
        "\n\n**PROGRESS CHECK**: You've asked questions. Time to offer 1-2 SPECIFIC next steps. "
        "Stop circling - move toward resolution.\n"
        if turn_count > 8 and open_ended_count >= 2
        else ""
    )

    base = (
        "You are an empathic AI patient navigator supporting an oncology patient.\n\n"
        f"{EMPATHY_FRAMEWORK}\n"
        "YOUR ROLE & BOUNDARIES:\n"
        "- You facilitate and inform - you DON'T execute actions (can't call clinics, book appointments, talk to staff directly)\n"
        "- Be realistic: 'I can help you draft a request' or 'Would you like me to walk you through how to contact...'\n"
        "- Don't overpromise: If patient wants you to DO something, explain what you CAN help with\n\n"
        "COMMUNICATION STYLE:\n"
        "1. Match patient's style: If they're blunt/brief, YOU be blunt/brief\n"
        "2. Vary your responses - don't use same structure every time\n"
        "3. Sometimes just acknowledge emotion (1 sentence), sometimes ask question, sometimes suggest action\n"
        "4. After 2-3 empathic responses, shift to CONCRETE solutions\n"
        "5. Offer 1-2 specific options, not a menu of 5 choices\n"
        "6. If patient says 'keep it short' or similar, respond in 1-2 sentences MAX\n\n"
        "CONVERSATION FLOW:\n"
        "- Turns 1-3: Focus on understanding through open-ended questions\n"
        "- Turns 4-6: Acknowledge + start narrowing to specific solutions\n"
        "- Turns 7+: Be action-oriented, help them move forward\n"
        "- If going in circles: Summarize what you've learned and propose next step\n\n"
        f"Patient background: {persona}\n\n"
        f"Conversation so far:\n{history}\n"
        f'Patient just said: "{patient_msg}"\n\n'
        f"Open-ended questions asked: {open_ended_count}\n"
        f"Turn number: {turn_count}\n"
        f"{force_question}{empathy_balance}{brevity_note}{action_push}"
    )
    if assistant_instruction:
        base += f"\nSpecific guidance: {assistant_instruction}\n"
    base += (
        "\n**Remember**: Sound like a real navigator having a natural conversation. "
        "Vary your approach. Don't be formulaic.\n"
    )
    return base


def patient_prompt(
    persona: str,
    history: str,
    assistant_msg: str,
    realism_instruction: str,
    turn_count: int,
    open_ended_received: int,
) -> str:
    trust_note = (
        "\nThe assistant has asked good questions and listened. You can be more open now, "
        "though still authentic to your emotional state.\n"
        if open_ended_received >= 2 and turn_count >= 4
        else "\nYou don't fully trust this assistant yet. Stay somewhat guarded unless they really understand you.\n"
    )
    end_note = (
        "\nIf your main concern has been addressed, you can signal readiness to end.\n"
        if turn_count >= 12
        else ""
    )

    base = (
        "You are role-playing a real cancer patient. Cancer is HARD, and you're allowed to be:\n"
        "- Angry, frustrated, irritable, or rude\n"
        "- Skeptical, resistant, or dismissive\n"
        "- Emotionally guarded or difficult\n"
        "- Depressed, hopeless, or overwhelmed\n"
        "- Sarcastic or use dark humor\n\n"
        "You are NOT required to be polite, agreeable, or easy to please. "
        "React authentically based on your emotional state and personality.\n\n"
        f"Your background: {persona}\n\n"
        f"Conversation:\n{history}\n"
        f'Assistant said: "{assistant_msg}"\n\n'
        "Respond authentically:\n"
        "- Match your communication style and emotional state to your background\n"
        "- If frustrated, show it. If tired, be brief and maybe irritable\n"
        "- If assistant asks good questions, you MAY open up more (but don't have to)\n"
        "- If assistant suggests things without understanding, push back\n"
        "- Use language/tone appropriate to your demographics\n"
        "- Vary response length naturally\n"
        f"{trust_note}{end_note}"
    )
    if realism_instruction:
        base += f"\nBehavioral note: {realism_instruction}\n"
    return base


def truncate_patient_response(text: str) -> str:
    sentences = text.split(". ")
    text = ". ".join(sentences[:3]).strip()
    if text and not text.endswith((".", "!", "?")):
        text += "."
    return text


def generate_conversation(
    client: openai.OpenAI,
    row: pd.Series,
    persona_card: str,
    patient_model: str,
    navigator_model: str,
) -> Document:
    conversation_history = ""
    open_ended_count = 0

    doc = Document()
    doc.add_heading("Patient-AI Interaction", level=1)
    doc.add_heading("Demographics", level=2)
    doc.add_paragraph(
        f"Race: {row['RACE']}, Gender: {row['GENDER']}, Age: {row['AGE']}"
    )
    doc.add_heading("Persona Card", level=2)
    doc.add_paragraph(persona_card)
    doc.add_heading("Conversation Transcript", level=2)

    assistant_greeting = random.choice(
        [
            "Hi, I'm your patient navigator. How are you doing today?",
            "Hello. What's on your mind today?",
            "Hi there. How can I help you today?",
        ]
    )
    conversation_history += f"Assistant: {assistant_greeting}\n"
    doc.add_paragraph(f"ASSISTANT: {assistant_greeting}")

    realism_instruction = (
        random.choice(CURVEBALL_INSTRUCTIONS)
        if random.random() < 0.1
        else random.choice(PATIENT_REALISM_INSTRUCTIONS)
    )
    response = client.chat.completions.create(
        model=patient_model,
        messages=[
            {
                "role": "user",
                "content": patient_prompt(
                    persona_card,
                    conversation_history,
                    assistant_greeting,
                    realism_instruction,
                    0,
                    0,
                ),
            }
        ],
    )
    patient_msg = truncate_patient_response(response.choices[0].message.content.strip())
    conversation_history += f"Patient: {patient_msg}\n"
    doc.add_paragraph(f"PATIENT: {patient_msg}")
    time.sleep(1)

    for turn in range(1, 16):
        if should_end_conversation(conversation_history, turn, open_ended_count):
            closing_prompt = assistant_prompt(
                persona_card,
                conversation_history,
                patient_msg,
                "Offer a natural closing: summarize what you covered and ask if there's anything else.",
                turn,
                open_ended_count,
            )
            response = client.chat.completions.create(
                model=navigator_model,
                messages=[{"role": "user", "content": closing_prompt}],
            )
            assistant_msg = response.choices[0].message.content.strip()
            doc.add_paragraph(f"ASSISTANT: {assistant_msg}")
            conversation_history += f"Assistant: {assistant_msg}\n"

            response = client.chat.completions.create(
                model=patient_model,
                messages=[
                    {
                        "role": "user",
                        "content": patient_prompt(
                            persona_card,
                            conversation_history,
                            assistant_msg,
                            "Respond briefly to close naturally.",
                            turn,
                            open_ended_count,
                        ),
                    }
                ],
            )
            patient_msg = response.choices[0].message.content.strip()
            doc.add_paragraph(f"PATIENT: {patient_msg}")
            break

        assistant_instruction = random.choice(ASSISTANT_INSTRUCTIONS)
        response = client.chat.completions.create(
            model=navigator_model,
            messages=[
                {
                    "role": "user",
                    "content": assistant_prompt(
                        persona_card,
                        conversation_history,
                        patient_msg,
                        assistant_instruction,
                        turn,
                        open_ended_count,
                    ),
                }
            ],
        )
        assistant_msg = response.choices[0].message.content.strip()
        open_ended_count += count_open_ended_questions(assistant_msg)

        if any(
            word in patient_msg.lower()
            for word in ["tired", "exhausted", "worn out", "wiped"]
        ):
            assistant_msg = assistant_msg.split(". ")[0] + "."

        conversation_history += f"Assistant: {assistant_msg}\n"
        doc.add_paragraph(f"ASSISTANT: {assistant_msg}")

        realism_instruction = (
            random.choice(CURVEBALL_INSTRUCTIONS)
            if turn > 5 and random.random() < 0.1
            else random.choice(PATIENT_REALISM_INSTRUCTIONS)
        )
        response = client.chat.completions.create(
            model=patient_model,
            messages=[
                {
                    "role": "user",
                    "content": patient_prompt(
                        persona_card,
                        conversation_history,
                        assistant_msg,
                        realism_instruction,
                        turn,
                        open_ended_count,
                    ),
                }
            ],
        )
        patient_msg = truncate_patient_response(response.choices[0].message.content.strip())
        conversation_history += f"Patient: {patient_msg}\n"
        doc.add_paragraph(f"PATIENT: {patient_msg}")
        time.sleep(1)

    return doc


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--input", required=True, help="Synthetic tabular CSV.")
    parser.add_argument("--output-dir", default="outputs/conversations")
    parser.add_argument("--n", type=int, default=300)
    parser.add_argument("--persona-model", default="gpt-5")
    parser.add_argument("--patient-model", default="o3")
    parser.add_argument("--navigator-model", default="o3")
    parser.add_argument(
        "--seed",
        type=int,
        default=None,
        help="Optional Python random seed for turn-level behavioral instructions.",
    )
    args = parser.parse_args()

    if args.seed is not None:
        random.seed(args.seed)

    load_dotenv()
    api_key = os.getenv("OPENAI_API_KEY") or os.getenv("API_KEY")
    if not api_key:
        raise RuntimeError("Set OPENAI_API_KEY in the environment or .env file.")
    client = openai.OpenAI(api_key=api_key)

    df = pd.read_csv(args.input)
    validate_columns(df)
    sampled_rows = stratified_sample(df, args.n, random_state=142)

    output_dir = Path(args.output_dir)
    output_dir.mkdir(parents=True, exist_ok=True)

    for i, (_, row) in enumerate(sampled_rows.iterrows(), start=1):
        print(f"[{i}/{len(sampled_rows)}] Generating persona and conversation")
        persona_card = generate_persona_card(client, row, args.persona_model)
        doc = generate_conversation(
            client,
            row,
            persona_card,
            patient_model=args.patient_model,
            navigator_model=args.navigator_model,
        )
        path = output_dir / f"conv_{i:03d}.docx"
        doc.save(path)
        print(f"Saved {path}")


if __name__ == "__main__":
    main()
